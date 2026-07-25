import re
import sys
from pathlib import Path

from docx import Document

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"


def _resolve_template(name: str) -> Path | None:
    print(f"[TEMPLATE-FORM] Looking for template: {name!r}", file=sys.stderr)
    print(f"[TEMPLATE-FORM] TEMPLATES_DIR: {TEMPLATES_DIR} (exists={TEMPLATES_DIR.exists()})", file=sys.stderr)
    if TEMPLATES_DIR.exists():
        print(f"[TEMPLATE-FORM] Files in dir: {list(TEMPLATES_DIR.iterdir())}", file=sys.stderr)
    p = TEMPLATES_DIR / name
    if p.exists():
        print(f"[TEMPLATE-FORM] Found exact match: {p}", file=sys.stderr)
        return p
    for ext in (".docx", ".xlsx", ".pptx"):
        p = TEMPLATES_DIR / f"{name}{ext}"
        if p.exists():
            print(f"[TEMPLATE-FORM] Found with ext: {p}", file=sys.stderr)
            return p
    print(f"[TEMPLATE-FORM] Template not found!", file=sys.stderr)
    return None


def _find_all(text: str) -> list[str]:
    return list(dict.fromkeys(re.findall(r"\{\{.+?\}\}", text)))


def generate_template_form(name: str) -> str:
    """Generate a markdown data-entry form from a template's placeholders.

    The output format is designed to be parseable by generate_from_form tool.
    """
    path = _resolve_template(name)
    if path is None:
        return f"# Data Entry Form\n\nTemplate `{name}` not found.\n"

    doc = Document(str(path))

    # Collect all unique placeholders, noting their locations
    all_phs: dict[str, list[str]] = {}
    dynamic_tables: list[dict] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        for ph in _find_all(text):
            all_phs.setdefault(ph, []).append(text[:60])

    for ti, table in enumerate(doc.tables):
        has_dynamic = False
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue
            for cell in row.cells:
                if _find_all(cell.text):
                    has_dynamic = True
                    break
            if has_dynamic:
                break

        if not has_dynamic:
            for ri, row in enumerate(table.rows):
                for cell in row.cells:
                    for ph in _find_all(cell.text):
                        all_phs.setdefault(ph, []).append(f"Table {ti+1}: {cell.text[:40]}")
        else:
            col_info: list[dict] = []
            for ci, cell in enumerate(table.rows[1].cells):
                phs = _find_all(cell.text)
                if phs:
                    col_info.append({"col": ci, "placeholder": phs[0], "name": re.sub(r"\{\{|\}\}", "", phs[0])})
            dynamic_tables.append({
                "index": ti,
                "header": [c.text[:30] for c in table.rows[0].cells],
                "cols": col_info,
            })

    lines = [f"# Data Entry Form: `{path.stem}`", ""]

    # Section 1: General Fields
    lines.append("## General Fields")
    lines.append("")
    lines.append("| Placeholder | Value |")
    lines.append("|-------------|-------|")
    for ph in sorted(all_phs.keys()):
        lines.append(f"| `{ph}` | |")
    lines.append("")

    # Section 2: Dynamic Tables
    if dynamic_tables:
        lines.append("## Dynamic Tables (fill_table_rows)")
        lines.append("")
        for t in dynamic_tables:
            ti = t["index"]
            hdr = " | ".join(t["header"])
            lines.append(f"### Table {ti + 1}")
            lines.append(f"Header: {hdr}")
            lines.append("")
            if t["cols"]:
                header_row = " | ".join(f"`{{{{{c['name']}}}}}`" for c in t["cols"])
                sep = " | ".join("---" for _ in t["cols"])
                lines.append(f"| {header_row} |")
                lines.append(f"| {sep} |")
                lines.append("")

    lines.append("---")
    lines.append(
        "Fill the values above and call `generate_from_form` to produce the document."
    )
    lines.append("")
    return "\n".join(lines)
