import json
import re
import os
from pathlib import Path
from typing import Any

from docx import Document
from office_mcp.core.client import run_officecli
from office_mcp.tools.tables import fill_table_rows

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", str(Path(__file__).parent.parent / "output")))


def _resolve_template(template: str) -> Path | None:
    p = TEMPLATES_DIR / template
    if p.exists():
        return p
    for ext in (".docx", ".xlsx", ".pptx"):
        p = TEMPLATES_DIR / f"{template}{ext}"
        if p.exists():
            return p
    return None


def _parse_general_fields(md: str) -> dict[str, str]:
    """Parse the General Fields table from the markdown form."""
    data = {}
    in_general = False
    for line in md.splitlines():
        if line.strip().startswith("## General Fields"):
            in_general = True
            continue
        if in_general:
            if line.strip().startswith("## "):
                break
            if line.strip().startswith("|") and "{{" in line and "|" in line:
                parts = [p.strip() for p in line.split("|")]
                if len(parts) >= 3:
                    key = parts[1].strip()
                    value = parts[2].strip()
                    key = key.replace("{{", "").replace("}}", "").replace("`", "").strip()
                    data[key] = value
    return data


def _strip_placeholder_brackets(text: str) -> str:
    return text.replace("{{", "").replace("}}", "").replace("`", "").strip()


def _parse_table_sections(md: str, doc: Document) -> dict[int, list[dict[str, str]]]:
    """Parse table data rows from the markdown form.

    Returns dict mapping 0-based table_index to list of row dicts.
    """
    sections: dict[int, list[str]] = {}
    current_section = None
    current_lines: list[str] = []

    for line in md.splitlines():
        m = re.match(r"^#{2,3} Table (\d+)", line.strip())
        if m:
            if current_section is not None and _has_table_rows(current_lines):
                sections[int(current_section)] = current_lines
            current_section = m.group(1)
            current_lines = [line]
        elif current_section is not None:
            current_lines.append(line)

    if current_section is not None and _has_table_rows(current_lines):
        sections[int(current_section)] = current_lines

    result: dict[int, list[dict[str, str]]] = {}

    for ti_1based, lines in sections.items():
        ti = ti_1based - 1
        if ti >= len(doc.tables):
            continue

        table = doc.tables[ti]
        raw_rows = _extract_table_rows(lines)
        if not raw_rows:
            continue

        # Map column position to placeholder name
        template_row_idx = None
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                if re.search(r"\{\{.+?\}\}", cell.text):
                    template_row_idx = ri
                    break
            if template_row_idx is not None:
                break

        if template_row_idx is None:
            continue

        # Get column placeholders from the template row
        col_phs: list[str | None] = []
        for cell in table.rows[template_row_idx].cells:
            phs = re.findall(r"\{\{(.+?)\}\}", cell.text)
            col_phs.append(phs[0] if phs else None)

        # Determine column order from the markdown header
        header_line = None
        for line in lines:
            if line.strip().startswith("|") and "|" in line:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if cells and all("{{" in c or c == "---" for c in cells[:3]):
                    header_line = cells
                    break

        if not header_line:
            continue

        # Map markdown columns to placeholder names
        col_map: list[str | None] = []
        for h in header_line:
            if "{{" in h:
                col_map.append(re.sub(r"\{\{|\}\}", "", h).strip())
            else:
                col_map.append(None)

        rows_data: list[dict[str, str]] = []
        in_data = False
        for line in lines:
            if line.strip().startswith("|") and "|" in line:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if not in_data:
                    if cells and all(c == "---" for c in cells):
                        in_data = True
                    continue
                if len(cells) < len(header_line):
                    continue
                if all(c == "" for c in cells):
                    continue
                row: dict[str, str] = {}
                for ci, col_name in enumerate(col_map):
                    if col_name and ci < len(cells):
                        row[col_name] = cells[ci].replace("\\n", "\n")
                if row:
                    rows_data.append(row)

        if rows_data:
            result[ti] = rows_data

    return result


def _has_table_rows(lines: list[str]) -> bool:
    """Check if the section has at least one data row (after the separator)."""
    found_sep = False
    for line in lines:
        cells = [c.strip() for c in line.split("|")[1:-1]] if line.strip().startswith("|") else []
        if cells and all(c == "---" for c in cells):
            found_sep = True
        elif found_sep and cells and any(c for c in cells):
            return True
    return False


def _extract_table_rows(lines: list[str]) -> list[list[str]]:
    """Extract raw cell values from table markdown."""
    rows: list[list[str]] = []
    in_data = False
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if not in_data:
            if cells and all(c == "---" for c in cells):
                in_data = True
            continue
        if all(c == "" for c in cells):
            continue
        rows.append(cells)
    return rows


def generate_from_form(
    template: str,
    form_content: str,
    output_filename: str | None = None,
) -> dict[str, Any]:
    """Generate a document from a filled markdown data-entry form.

    The form must follow the format produced by template://{name}/form.
    The function parses general fields and table rows, then generates
    the document via create_doc + fill_table_rows.

    Args:
        template: Template name or filename
        form_content: The filled markdown form content
        output_filename: Name for the generated file (optional, auto-generated if omitted)
    """
    template_path = _resolve_template(template)
    if template_path is None:
        return {"success": False, "error": f"Template '{template}' not found"}

    if not output_filename:
        output_filename = f"{template_path.stem}_output.docx"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / output_filename

    doc = Document(str(template_path))

    # Parse general fields
    general_data = _parse_general_fields(form_content)

    # Parse table rows
    table_rows = _parse_table_sections(form_content, doc)

    # Step 1: Create the initial document (merge with only general fields)
    if general_data:
        result = run_officecli(
            "merge",
            str(template_path),
            str(output_path),
            "--data",
            json.dumps(general_data),
            timeout=60,
        )
        if result.get("error"):
            return {"success": False, "error": f"Merge failed: {result['error']}"}

        file_to_edit = str(output_path)
    else:
        import shutil
        shutil.copy2(str(template_path), str(output_path))
        file_to_edit = str(output_path)

    # Step 2: Fill table rows for each dynamic table
    filled_tables = []
    for ti in sorted(table_rows.keys()):
        rows = table_rows[ti]
        if not rows:
            continue

        # The fill_table_rows tool needs the placeholder keys to match
        # what's in the document. The keys are already the placeholder names
        # (without {{}}).
        tr = fill_table_rows(
            filepath=file_to_edit,
            rows=rows,
            table_index=ti,
        )
        if not tr.get("success"):
            return {
                "success": False,
                "error": f"Failed to fill table {ti + 1}: {tr.get('error')}",
            }
        filled_tables.append({"table_index": ti, "rows": len(rows)})

    return {
        "success": True,
        "filename": output_filename,
        "path": str(output_path),
        "format": "docx",
        "template": template,
        "general_fields": len(general_data),
        "tables_filled": filled_tables,
    }
