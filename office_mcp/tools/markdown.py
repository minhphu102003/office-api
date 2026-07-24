import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

H1_SIZE = Pt(18)
H2_SIZE = Pt(16)
H3_SIZE = Pt(14)
BODY_SIZE = Pt(12)
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
BLACK = RGBColor(0, 0, 0)
FONT_NAME = "Times New Roman"


def _setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    rpr = normal.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
            f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
            f"</w:rPr>"
        )
        normal.element.append(rpr)

    for level, size, before in [
        ("Heading 1", H1_SIZE, Pt(18)),
        ("Heading 2", H2_SIZE, Pt(12)),
        ("Heading 3", H3_SIZE, Pt(12)),
    ]:
        s = doc.styles[level]
        s.font.name = FONT_NAME
        s.font.size = size
        s.font.bold = True
        s.font.color.rgb = DARK_BLUE
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.space_before = before
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5
        rpr = s.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = parse_xml(
                f'<w:rPr {nsdecls("w")}>'
                f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
                f"</w:rPr>"
            )
            s.element.append(rpr)


def _add_formatted_text(paragraph, text: str):
    pattern = r"(\*\*.*?\*\*|\*.*?\*|{{.*?}})"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = FONT_NAME
        elif part.startswith("*") and not part.startswith("**") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = FONT_NAME
        else:
            run = paragraph.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE


def _parse_markdown(md: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    in_list = False
    for line in md.split("\n"):
        stripped = line.strip()
        if not stripped:
            if in_list:
                blocks.append({"type": "list_gap"})
                in_list = False
            else:
                blocks.append({"type": "paragraph", "text": ""})
            continue

        h = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if h:
            in_list = False
            blocks.append({"type": f"h{len(h.group(1))}", "text": h.group(2)})
            continue

        if re.match(r"^[-*+]\s+", stripped):
            in_list = True
            text = re.sub(r"^[-*+]\s+", "", stripped)
            blocks.append({"type": "bullet", "text": text})
            continue

        if re.match(r"^\d+[.)]\s+", stripped):
            in_list = True
            text = re.sub(r"^\d+[.)]\s+", "", stripped)
            blocks.append({"type": "ordered", "text": text, "index": int(re.match(r"\d+", stripped).group())})
            continue

        if stripped.startswith(">"):
            in_list = False
            blocks.append({"type": "quote", "text": re.sub(r"^>\s?", "", stripped)})
            continue

        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            in_list = False
            blocks.append({"type": "page_break"})
            continue

        in_list = False
        blocks.append({"type": "paragraph", "text": stripped})
    return blocks


def markdown_to_template(
    markdown: str,
    output_filename: str,
    title: str | None = None,
) -> dict[str, Any]:
    """Convert markdown content into a formatted .docx template. Handles headings, paragraphs, bold/italic, bullet lists, numbered lists, blockquotes, and {{placeholder}} markers. Formatting follows the word-format skill standards (Times New Roman, 1.5 spacing, justified body, dark blue headings).

    Args:
        markdown: Markdown content with optional {{placeholder}} markers
        output_filename: Template filename with .docx extension (e.g. 'contract_template.docx')
        title: Optional document title (appears as centered Heading 1 at the top)
    """
    if not output_filename.lower().endswith(".docx"):
        return {"success": False, "error": "Only .docx is supported for markdown conversion"}

    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    _setup_styles(doc)

    if title:
        p = doc.add_heading(title, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    blocks = _parse_markdown(markdown)

    i = 0
    while i < len(blocks):
        block = blocks[i]
        t = block["type"]

        if t == "h1":
            doc.add_heading(block["text"], level=1)
        elif t == "h2":
            doc.add_heading(block["text"], level=2)
        elif t == "h3":
            doc.add_heading(block["text"], level=3)
        elif t == "paragraph":
            p = doc.add_paragraph()
            _add_formatted_text(p, block["text"])
        elif t == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, block["text"])
        elif t == "ordered":
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, block["text"])
        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block["text"])
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif t == "page_break":
            doc.add_page_break()
        elif t == "list_gap":
            if i + 1 < len(blocks) and blocks[i + 1]["type"] in ("bullet", "ordered"):
                doc.add_paragraph()

        i += 1

    filepath = TEMPLATES_DIR / output_filename
    doc.save(str(filepath))
    return {
        "success": True,
        "name": Path(output_filename).stem,
        "filename": output_filename,
        "format": "docx",
        "size": filepath.stat().st_size,
    }
