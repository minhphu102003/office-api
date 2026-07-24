import copy
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", str(Path(__file__).parent.parent / "output")))


def _has_placeholder(text: str) -> bool:
    return bool(re.search(r'\{\{.+?\}\}', text))


def _replace_in_xml(element, old_text: str, new_text: str):
    for r_elem in element.iter(qn("w:r")):
        for t_elem in r_elem.findall(qn("w:t")):
            if t_elem.text and old_text in t_elem.text:
                t_elem.text = t_elem.text.replace(old_text, new_text)


def _fix_newlines(element):
    """Convert literal \\n in <w:t> text into <w:br/> line breaks."""
    for t_elem in list(element.iter(qn("w:t"))):
        if t_elem.text and "\n" in t_elem.text:
            lines = t_elem.text.split("\n")
            t_elem.text = lines[0]
            for line in reversed(lines[1:]):
                if line:
                    new_t = copy.deepcopy(t_elem)
                    new_t.text = line
                    t_elem.addnext(new_t)
                br = OxmlElement("w:br")
                t_elem.addnext(br)


def _resolve_path(filepath: str) -> Path | None:
    for base in (TEMPLATES_DIR, OUTPUT_DIR, Path(".").resolve()):
        candidate = base / filepath
        if candidate.exists():
            return candidate.resolve()
        for ext in (".docx", ".xlsx", ".pptx"):
            candidate = base / f"{filepath}{ext}"
            if candidate.exists():
                return candidate.resolve()
    p = Path(filepath)
    if p.exists():
        return p.resolve()
    return None


def fill_table_rows(
    filepath: str,
    rows: list[dict[str, str]],
    table_index: int = 0,
) -> dict[str, Any]:
    """Clone the template row in a table N times and fill placeholders.

    The document must contain a table with at least one data row containing
    {{placeholder}} markers (e.g. {{project_name}}, {{progress_percent}}).
    That row will be cloned for each item in `rows`, then removed after cloning.

    Works on both template files (in the templates dir) and generated documents
    (in the output dir).

    Args:
        filepath: Filename or path to the .docx file
        rows: List of dicts mapping placeholder names to values, one per output row
        table_index: 0-based index of the target table (default 0)
    """
    resolved = _resolve_path(filepath)
    if resolved is None:
        return {"success": False, "error": f"File '{filepath}' not found in templates or output directory"}

    ext = resolved.suffix.lower()
    if ext != ".docx":
        return {"success": False, "error": f"Unsupported format '{ext}'. Only .docx is supported"}

    try:
        doc = Document(str(resolved))
    except Exception as e:
        return {"success": False, "error": f"Failed to open document: {e}"}

    if table_index >= len(doc.tables):
        return {
            "success": False,
            "error": f"Table index {table_index} not found (document has {len(doc.tables)} tables)",
        }

    table = doc.tables[table_index]

    # Find the first data row that contains {{...}} placeholders
    template_row_idx = None
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if _has_placeholder(cell.text):
                template_row_idx = i
                break
        if template_row_idx is not None:
            break

    if template_row_idx is None:
        return {"success": False, "error": "No template row found — no row contains {{...}} placeholders"}

    template_tr = table.rows[template_row_idx]._tr

    # Clone the template row for each data item
    for item in rows:
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)

        for key, value in item.items():
            placeholder = "{{" + key + "}}"
            _replace_in_xml(new_tr, placeholder, value)

        _fix_newlines(new_tr)

    # Remove the original template row
    table._tbl.remove(template_tr)

    try:
        doc.save(str(resolved))
    except Exception as e:
        return {"success": False, "error": f"Failed to save document: {e}"}

    return {
        "success": True,
        "filename": resolved.name,
        "path": str(resolved),
        "table_index": table_index,
        "rows_filled": len(rows),
        "template_row_removed": True,
    }
